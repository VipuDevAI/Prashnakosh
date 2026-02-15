#!/usr/bin/env python3
"""
Enhanced Word Document Parser for Question Banks and Question Papers
Handles various formats including embedded page markers
"""

import json
import re
from docx import Document
from typing import List, Dict, Any, Optional

def clean_text(text: str) -> str:
    """Clean and normalize text"""
    # Remove page markers (various formats)
    text = re.sub(r'Page:\s*\d+/\d+', ' ', text)
    text = re.sub(r'Page\s+\d+\s+of\s+\d+', ' ', text)
    # Normalize whitespace
    text = ' '.join(text.split())
    return text.strip()

def extract_marks(text: str) -> Optional[int]:
    """Extract marks from question text"""
    patterns = [
        r'\[(\d+)\s*(?:mark|marks|m)\]',
        r'\((\d+)\s*(?:mark|marks|m)\)',
        r'(\d+)\s*(?:mark|marks)\s*$',
        r'\[(\d+)\]',
        r'Marks\s*(\d+)',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    return None

def detect_question_type(text: str, options: List[str]) -> str:
    """Detect question type based on content"""
    text_lower = text.lower()
    
    if options and len(options) >= 2:
        return "mcq"
    if "true or false" in text_lower or "true/false" in text_lower:
        return "true_false"
    if "assertion" in text_lower and "reason" in text_lower:
        return "assertion_reason"
    if "fill in" in text_lower or "fill up" in text_lower:
        return "fill_blank"
    if "match" in text_lower and "column" in text_lower:
        return "matching"
    
    # Based on marks
    marks = extract_marks(text)
    if marks:
        if marks >= 4:
            return "long_answer"
        elif marks >= 2:
            return "short_answer"
    
    return "short_answer"

def parse_document_universal(filepath: str, source_name: str, chapter_default: str = "Mixed") -> List[Dict[str, Any]]:
    """Universal parser that handles multiple document formats"""
    doc = Document(filepath)
    questions = []
    
    # Combine all text
    full_text = ""
    for para in doc.paragraphs:
        full_text += para.text + "\n"
    
    # Clean page markers
    full_text = re.sub(r'Page:\s*\d+/\d+', '\n', full_text)
    full_text = re.sub(r'Page\s+\d+\s+of\s+\d+', '\n', full_text)
    
    # Split into lines and clean
    lines = [clean_text(l) for l in full_text.split('\n')]
    lines = [l for l in lines if l]
    
    current_question = None
    current_q_num = None
    current_options = []
    current_marks = None
    question_buffer = []
    
    # Patterns for question detection
    q_patterns = [
        r'^Q\.?\s*No\.?\s*(\d+)',  # Q No. 1
        r'^Q\.?\s*(\d+)[.\):\s]',   # Q.1 or Q1) or Q1:
        r'^(\d+)[.\)]\s+',          # 1. or 1)
        r'^(\d+)\s+(?=[A-Z])',      # 1 What...
    ]
    
    option_pattern = r'^[(\[]?([a-dA-D])[)\].\s]+(.+)'
    
    def save_current_question():
        nonlocal current_question, current_options, current_marks, question_buffer
        
        if question_buffer:
            q_text = ' '.join(question_buffer)
            q_text = re.sub(r'\s*(Marks?|marks?)\s*\d*\s*$', '', q_text).strip()
            q_text = re.sub(r'\[?\d+\]?\s*$', '', q_text).strip()
            
            if len(q_text) > 15:  # Minimum question length
                q_type = detect_question_type(q_text, current_options)
                if not current_marks:
                    current_marks = 1 if q_type == "mcq" else 2
                
                questions.append({
                    "questionText": q_text,
                    "type": q_type,
                    "marks": current_marks,
                    "options": current_options if current_options else None,
                    "correctAnswer": None,
                    "chapter": chapter_default,
                    "source": source_name,
                })
        
        question_buffer = []
        current_options = []
        current_marks = None
    
    for line in lines:
        # Skip headers and instructions
        if re.match(r'^(Section|SECTION|General Instructions|Time Allowed|Maximum Marks|CLASS|COMPUTER SCIENCE|MARKING SCHEME|Q\s*No\s+Section)', line, re.IGNORECASE):
            continue
        
        # Check for new question
        new_q = False
        for pattern in q_patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                save_current_question()
                new_q = True
                current_q_num = match.group(1)
                # Get rest of line after question number
                rest = re.sub(pattern, '', line, flags=re.IGNORECASE).strip()
                if rest:
                    question_buffer.append(rest)
                current_marks = extract_marks(line)
                break
        
        if new_q:
            continue
        
        # Check for options
        opt_match = re.match(option_pattern, line)
        if opt_match and question_buffer:
            option_text = f"{opt_match.group(1).upper()}) {opt_match.group(2)}"
            current_options.append(option_text)
            continue
        
        # Skip OR markers
        if re.match(r'^OR\s*$', line, re.IGNORECASE):
            continue
        
        # Add to question buffer
        if question_buffer or current_q_num:
            question_buffer.append(line)
    
    # Don't forget last question
    save_current_question()
    
    return questions

def parse_chapter_bank(filepath: str, chapter_name: str) -> List[Dict[str, Any]]:
    """Parse chapter-wise question bank with special handling"""
    doc = Document(filepath)
    questions = []
    
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Look for structured Q&A patterns
    # Pattern: Q1. question text... Ans: answer
    
    q_pattern = r'(?:Q\.?\s*)?(\d+)[.\)]\s*(.+?)(?=(?:Q\.?\s*)?\d+[.\)]|$)'
    matches = re.findall(q_pattern, full_text, re.DOTALL | re.IGNORECASE)
    
    for q_num, q_content in matches:
        q_content = clean_text(q_content)
        
        if len(q_content) < 15:
            continue
        
        # Extract answer if present
        answer = None
        ans_match = re.search(r'(?:Ans(?:wer)?|Correct\s*Answer)[:\s]+(.+?)(?=\s*(?:Q\.?\s*)?\d+[.\)]|$)', q_content, re.IGNORECASE)
        if ans_match:
            answer = clean_text(ans_match.group(1))
            q_content = q_content[:ans_match.start()].strip()
        
        # Extract options
        options = []
        for opt_match in re.finditer(r'[(\[]?([a-dA-D])[)\].\s]+([^(\[a-dA-D]+?)(?=[(\[]?[a-dA-D][)\].\s]|$)', q_content):
            options.append(f"{opt_match.group(1).upper()}) {clean_text(opt_match.group(2))}")
        
        # Clean question text (remove options if found inline)
        if options:
            first_opt = re.search(r'[(\[]?[a-dA-D][)\].\s]', q_content)
            if first_opt:
                q_content = q_content[:first_opt.start()].strip()
        
        q_type = detect_question_type(q_content, options)
        marks = extract_marks(q_content) or (1 if q_type == "mcq" else 2)
        
        # Clean marks from question text
        q_content = re.sub(r'\s*\[?\d+\s*(?:mark|marks|m)?\]?\s*$', '', q_content, flags=re.IGNORECASE).strip()
        
        if len(q_content) > 15:
            questions.append({
                "questionText": q_content,
                "type": q_type,
                "marks": marks,
                "options": options if options else None,
                "correctAnswer": answer,
                "chapter": chapter_name,
                "source": f"Chapter Bank - {chapter_name}",
            })
    
    return questions

def main():
    all_questions = []
    
    # Parse Chapter Question Banks
    chapter_files = {
        "question_docs/chapter2.docx": "Functions",
        "question_docs/chapter4.docx": "File Handling", 
        "question_docs/chapter5.docx": "Exception Handling",
    }
    
    print("=" * 60)
    print("PARSING CHAPTER-WISE QUESTION BANKS")
    print("=" * 60)
    
    for filepath, chapter in chapter_files.items():
        try:
            questions = parse_chapter_bank(filepath, chapter)
            if len(questions) < 5:  # Fallback to universal parser
                questions = parse_document_universal(filepath, f"Chapter - {chapter}", chapter)
            all_questions.extend(questions)
            print(f"✓ {chapter}: {len(questions)} questions parsed")
        except Exception as e:
            print(f"✗ Error parsing {filepath}: {e}")
    
    # Parse Question Papers with universal parser
    paper_files = {
        "question_docs/ssm_final.docx": "SSM Final 2025-26",
        "question_docs/kv_region.docx": "KV Region Preboard 2025-26",
        "question_docs/qp_xii_1.docx": "QP XII-1",
        "question_docs/sqp_24_25.docx": "SQP 2024-25",
    }
    
    print("\n" + "=" * 60)
    print("PARSING FULL QUESTION PAPERS")
    print("=" * 60)
    
    for filepath, paper_name in paper_files.items():
        try:
            questions = parse_document_universal(filepath, paper_name)
            all_questions.extend(questions)
            print(f"✓ {paper_name}: {len(questions)} questions parsed")
        except Exception as e:
            print(f"✗ Error parsing {filepath}: {e}")
    
    # Remove duplicates based on question text similarity
    seen = set()
    unique_questions = []
    for q in all_questions:
        q_hash = q["questionText"][:100].lower().strip()
        if q_hash not in seen:
            seen.add(q_hash)
            unique_questions.append(q)
    
    # Summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total Questions Parsed: {len(all_questions)}")
    print(f"Unique Questions: {len(unique_questions)}")
    
    # Count by type
    type_counts = {}
    for q in unique_questions:
        t = q.get("type", "unknown")
        type_counts[t] = type_counts.get(t, 0) + 1
    
    print("\nBy Type:")
    for t, count in sorted(type_counts.items()):
        print(f"  - {t}: {count}")
    
    # Count by chapter/source
    source_counts = {}
    for q in unique_questions:
        src = q.get("source", "Unknown")
        source_counts[src] = source_counts.get(src, 0) + 1
    
    print("\nBy Source:")
    for src, count in sorted(source_counts.items()):
        print(f"  - {src}: {count}")
    
    # Save to JSON
    output_file = "parsed_questions.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(unique_questions, f, indent=2, ensure_ascii=False)
    
    print(f"\n✓ Saved {len(unique_questions)} unique questions to {output_file}")
    
    # Sample questions for verification
    print("\n" + "=" * 60)
    print("SAMPLE QUESTIONS (first 3)")
    print("=" * 60)
    for i, q in enumerate(unique_questions[:3]):
        print(f"\n--- Question {i+1} ---")
        print(f"Type: {q['type']}, Marks: {q['marks']}")
        print(f"Source: {q['source']}")
        print(f"Text: {q['questionText'][:150]}...")
        if q.get('options'):
            print(f"Options: {q['options'][:2]}...")
    
    return unique_questions

if __name__ == "__main__":
    main()
