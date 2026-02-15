#!/usr/bin/env python3
"""
Comprehensive Word Document Parser for Question Banks and Question Papers
Handles both chapter-wise question banks and full question papers
"""

import json
import re
import sys
from docx import Document
from typing import List, Dict, Any, Optional
import uuid

def clean_text(text: str) -> str:
    """Clean and normalize text"""
    # Remove page markers
    text = re.sub(r'Page:\s*\d+/\d+', '', text)
    # Normalize whitespace
    text = ' '.join(text.split())
    return text.strip()

def extract_marks(text: str) -> Optional[int]:
    """Extract marks from question text"""
    patterns = [
        r'\[(\d+)\s*(?:mark|marks|m)\]',
        r'\((\d+)\s*(?:mark|marks|m)\)',
        r'(\d+)\s*(?:mark|marks)\s*$',
        r'\[(\d+)\]',
    ]
    for pattern in patterns:
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            return int(match.group(1))
    return None

def detect_question_type(text: str, options: List[str]) -> str:
    """Detect question type based on content"""
    text_lower = text.lower()
    
    if options and len(options) >= 2:
        return "mcq"
    if "true or false" in text_lower or "true/false" in text_lower:
        return "true_false"
    if "assertion" in text_lower and "reason" in text_lower:
        return "assertion_reason"
    if "fill in" in text_lower or "fill up" in text_lower:
        return "fill_blank"
    if "match" in text_lower and "column" in text_lower:
        return "matching"
    
    # Based on marks
    marks = extract_marks(text)
    if marks:
        if marks >= 4:
            return "long_answer"
        elif marks >= 2:
            return "short_answer"
    
    return "short_answer"

def parse_chapter_question_bank(filepath: str, chapter_name: str) -> List[Dict[str, Any]]:
    """Parse chapter-wise question bank documents"""
    doc = Document(filepath)
    questions = []
    current_question = None
    current_options = []
    current_answer = None
    
    full_text = "\n".join([para.text for para in doc.paragraphs])
    
    # Split by question patterns
    lines = full_text.split('\n')
    
    for line in lines:
        line = clean_text(line)
        if not line:
            continue
        
        # Check for question start
        q_match = re.match(r'^(?:Q\.?\s*)?(\d+)[.\)]\s*(.+)', line)
        if q_match:
            # Save previous question
            if current_question:
                q_type = detect_question_type(current_question, current_options)
                marks = extract_marks(current_question) or (1 if q_type == "mcq" else 2)
                questions.append({
                    "questionText": re.sub(r'\[?\d+\s*(?:mark|marks|m)?\]?$', '', current_question, flags=re.IGNORECASE).strip(),
                    "type": q_type,
                    "marks": marks,
                    "options": current_options if current_options else None,
                    "correctAnswer": current_answer,
                    "chapter": chapter_name,
                })
            
            current_question = q_match.group(2)
            current_options = []
            current_answer = None
            continue
        
        # Check for options
        opt_match = re.match(r'^[(\[]?([a-dA-D])[)\].\s]+(.+)', line)
        if opt_match and current_question:
            current_options.append(f"{opt_match.group(1)}) {opt_match.group(2)}")
            continue
        
        # Check for answer
        ans_match = re.match(r'^(?:Ans(?:wer)?|Correct\s*Answer)[:\s]+(.+)', line, re.IGNORECASE)
        if ans_match:
            current_answer = ans_match.group(1).strip()
            continue
        
        # Append to current question if exists
        if current_question and line:
            current_question += " " + line
    
    # Don't forget last question
    if current_question:
        q_type = detect_question_type(current_question, current_options)
        marks = extract_marks(current_question) or (1 if q_type == "mcq" else 2)
        questions.append({
            "questionText": re.sub(r'\[?\d+\s*(?:mark|marks|m)?\]?$', '', current_question, flags=re.IGNORECASE).strip(),
            "type": q_type,
            "marks": marks,
            "options": current_options if current_options else None,
            "correctAnswer": current_answer,
            "chapter": chapter_name,
        })
    
    return questions

def parse_question_paper(filepath: str, paper_name: str) -> List[Dict[str, Any]]:
    """Parse full question paper documents"""
    doc = Document(filepath)
    questions = []
    
    full_text = "\n".join([para.text for para in doc.paragraphs])
    # Clean page markers
    full_text = re.sub(r'Page:\s*\d+/\d+', '\n', full_text)
    
    lines = full_text.split('\n')
    
    current_question = None
    current_options = []
    current_marks = None
    in_options = False
    
    for i, line in enumerate(lines):
        line = clean_text(line)
        if not line:
            continue
        
        # Skip headers and section markers
        if re.match(r'^(Section|SECTION|General Instructions|Time|Maximum Marks|CLASS|COMPUTER SCIENCE)', line):
            continue
        
        # Check for question start (various formats)
        q_match = re.match(r'^(\d+)[.\)]\s*(.+)', line)
        if q_match:
            # Save previous question
            if current_question and len(current_question) > 10:
                q_type = detect_question_type(current_question, current_options)
                if not current_marks:
                    current_marks = 1 if q_type == "mcq" else 2
                questions.append({
                    "questionText": current_question,
                    "type": q_type,
                    "marks": current_marks,
                    "options": current_options if current_options else None,
                    "correctAnswer": None,
                    "chapter": "Mixed",
                    "source": paper_name,
                })
            
            current_question = q_match.group(2)
            current_options = []
            current_marks = extract_marks(line)
            in_options = False
            continue
        
        # Check for options (a), (b), (c), (d) or A), B), C), D)
        opt_match = re.match(r'^[(\[]?([a-dA-D])[)\].\s]+(.+)', line)
        if opt_match and current_question:
            option_text = f"{opt_match.group(1).upper()}) {opt_match.group(2)}"
            current_options.append(option_text)
            in_options = True
            continue
        
        # Check for OR questions
        if re.match(r'^OR\s*$', line, re.IGNORECASE):
            continue
        
        # Append to current question
        if current_question and not in_options:
            current_question += " " + line
    
    # Last question
    if current_question and len(current_question) > 10:
        q_type = detect_question_type(current_question, current_options)
        if not current_marks:
            current_marks = 1 if q_type == "mcq" else 2
        questions.append({
            "questionText": current_question,
            "type": q_type,
            "marks": current_marks,
            "options": current_options if current_options else None,
            "correctAnswer": None,
            "chapter": "Mixed",
            "source": paper_name,
        })
    
    return questions

def main():
    all_questions = []
    
    # Parse Chapter Question Banks
    chapter_files = {
        "question_docs/chapter2.docx": "Functions",
        "question_docs/chapter4.docx": "File Handling",
        "question_docs/chapter5.docx": "Exception Handling",
    }
    
    print("=" * 60)
    print("PARSING CHAPTER-WISE QUESTION BANKS")
    print("=" * 60)
    
    for filepath, chapter in chapter_files.items():
        try:
            questions = parse_chapter_question_bank(filepath, chapter)
            all_questions.extend(questions)
            print(f"✓ {chapter}: {len(questions)} questions parsed")
        except Exception as e:
            print(f"✗ Error parsing {filepath}: {e}")
    
    # Parse Question Papers
    paper_files = {
        "question_docs/ssm_final.docx": "SSM Final 2025-26",
        "question_docs/kv_region.docx": "KV Region Preboard 2025-26",
        "question_docs/qp_xii_1.docx": "QP XII-1",
        "question_docs/sqp_24_25.docx": "SQP 2024-25",
    }
    
    print("\n" + "=" * 60)
    print("PARSING FULL QUESTION PAPERS")
    print("=" * 60)
    
    for filepath, paper_name in paper_files.items():
        try:
            questions = parse_question_paper(filepath, paper_name)
            all_questions.extend(questions)
            print(f"✓ {paper_name}: {len(questions)} questions parsed")
        except Exception as e:
            print(f"✗ Error parsing {filepath}: {e}")
    
    # Summary
    print("\n" + "=" * 60)
    print("SUMMARY")
    print("=" * 60)
    print(f"Total Questions Parsed: {len(all_questions)}")
    
    # Count by type
    type_counts = {}
    for q in all_questions:
        t = q.get("type", "unknown")
        type_counts[t] = type_counts.get(t, 0) + 1
    
    print("\nBy Type:")
    for t, count in sorted(type_counts.items()):
        print(f"  - {t}: {count}")
    
    # Count by chapter/source
    chapter_counts = {}
    for q in all_questions:
        ch = q.get("chapter") or q.get("source", "Unknown")
        chapter_counts[ch] = chapter_counts.get(ch, 0) + 1
    
    print("\nBy Chapter/Source:")
    for ch, count in sorted(chapter_counts.items()):
        print(f"  - {ch}: {count}")
    
    # Save to JSON
    output_file = "parsed_questions.json"
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(all_questions, f, indent=2, ensure_ascii=False)
    
    print(f"\n✓ Saved to {output_file}")
    
    return all_questions

if __name__ == "__main__":
    main()
